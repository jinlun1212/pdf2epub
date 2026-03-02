#!/usr/bin/env python3
"""
Convert a chapter range from PDF into a hybrid reflowable EPUB.

Improvements vs initial version:
- Line-level extraction to reduce merged paragraphs.
- Styled inline spans (italic / bold / relative size classes).
- Figure/table region detection that expands around caption and removes figure-internal text.
- No full-page image fallback.
- Internal validation report + basic auto-tuning attempts.
"""

from __future__ import annotations

import argparse
import html
import json
import re
import shutil
import statistics
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

try:
    import fitz  # PyMuPDF
except ImportError as exc:
    raise SystemExit(
        "Missing dependency 'pymupdf'. Install with: python3 -m pip install pymupdf"
    ) from exc

try:
    from ebooklib import epub
except ImportError as exc:
    raise SystemExit(
        "Missing dependency 'ebooklib'. Install with: python3 -m pip install ebooklib"
    ) from exc

try:
    from PIL import Image
except ImportError:
    Image = None  # type: ignore[assignment]

try:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    DocxDocument = None  # type: ignore[assignment]
    WD_ALIGN_PARAGRAPH = None  # type: ignore[assignment]
    OxmlElement = None  # type: ignore[assignment]
    qn = None  # type: ignore[assignment]


CAPTION_RE = re.compile(r"^\s*(?:figure|table)\s*\d+\.\d+\b", re.IGNORECASE)
SECTION_RE = re.compile(r"^\s*(?:[1-3]?\d)\.(?:\d{1,2})\s+[A-Z].*")
EQUATION_SYMBOL_RE = re.compile(r"[=+\-*/<>]|(?:\bsqrt\b)|[∑∫√σμΔπ]")
SENTENCE_END_RE = re.compile(r"[.!?][\"')\]]?$")


DEFAULTS: Dict[str, object] = {
    "reader_profile": "apple_books",
    "dpi": 190,
    "thresholds.heading_min_font_size": 13.0,
    "thresholds.equation_symbol_density": 0.045,
    "thresholds.equation_max_chars": 140,
    "thresholds.figure_y_extension_ratio": 0.34,
    "thresholds.figure_x_margin_ratio": 0.02,
    "thresholds.min_visual_area_ratio": 0.0015,
    "thresholds.paragraph_gap_multiplier": 1.12,
    "epub_build_tag": "r5",
}


@dataclass
class SpanRun:
    text: str
    size: float
    bold: bool
    italic: bool


@dataclass
class LineBlock:
    text: str
    rect: fitz.Rect
    spans: List[SpanRun]
    max_font: float
    avg_font: float
    is_bold: bool
    is_italic: bool


@dataclass
class VisualBlock:
    rect: fitz.Rect
    source: str


@dataclass
class ContentBlock:
    kind: str
    rect: fitz.Rect
    text: str = ""
    page_num: int = 0
    lines: List[LineBlock] = field(default_factory=list)
    fig_id: Optional[str] = None
    img_name: Optional[str] = None
    figure_has_inline_caption: bool = False
    anchor_id: Optional[str] = None
    css_class: str = ""


@dataclass
class BuildResult:
    blocks: List[ContentBlock]
    images_written: int
    validation: Dict[str, object]
    page_stats: Dict[int, Dict[str, int]] = field(default_factory=dict)


REF_RE = re.compile(r"\b(?:Chapter\s+\d+|Figure\s+\d+\.\d+|Table\s+\d+\.\d+)\b", re.IGNORECASE)


def slugify(s: str) -> str:
    x = re.sub(r"[^A-Za-z0-9]+", "-", s.strip().lower()).strip("-")
    return x or "ref"


def load_simple_yaml(path: Path) -> Dict[str, object]:
    result: Dict[str, object] = {}
    if not path.exists():
        return result

    section: Optional[str] = None
    with path.open("r", encoding="utf-8") as fh:
        for raw in fh:
            line = raw.rstrip("\n")
            if not line.strip() or line.lstrip().startswith("#"):
                continue

            if line.startswith("  ") and section:
                inner = line.strip()
                if ":" not in inner:
                    continue
                k, v = inner.split(":", 1)
                result[f"{section}.{k.strip()}"] = coerce_value(v.strip())
                continue

            if ":" not in line:
                continue
            k, v = line.split(":", 1)
            key = k.strip()
            value = v.strip()
            if value == "":
                section = key
            else:
                section = None
                result[key] = coerce_value(value)

    return result


def coerce_value(v: str) -> object:
    if v.startswith('"') and v.endswith('"'):
        return v[1:-1]
    if v.startswith("'") and v.endswith("'"):
        return v[1:-1]
    if v.lower() in {"true", "false"}:
        return v.lower() == "true"
    if re.fullmatch(r"-?\d+", v):
        return int(v)
    if re.fullmatch(r"-?\d+\.\d+", v):
        return float(v)
    return v


def get_opt(cfg: Dict[str, object], key: str) -> object:
    return cfg[key] if key in cfg else DEFAULTS[key]


def normalize_text(s: str) -> str:
    s = s.replace("\xa0", " ")
    s = s.replace("\ufffd", "")
    s = re.sub(r"[\u0000-\u0008\u000B\u000C\u000E-\u001F]", " ", s)
    s = re.sub(r"[\u200B-\u200D\uFEFF]", "", s)
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def symbol_density(text: str) -> float:
    if not text:
        return 0.0
    cnt = len(EQUATION_SYMBOL_RE.findall(text))
    return cnt / max(len(text), 1)


def escape_attr(v: str) -> str:
    return html.escape(v, quote=True)


def is_header_footer(line: LineBlock, page_h: float) -> bool:
    t = line.text.strip()
    if not t:
        return True
    if re.match(r"^\s*CHAPTER\s+\d+\s*$", t, re.I):
        return False
    near_top = line.rect.y0 < page_h * 0.045
    near_bottom = line.rect.y1 > page_h * 0.955
    short = len(t) < 42
    if (near_top or near_bottom) and short:
        return True
    if re.fullmatch(r"\d+", t):
        return True
    if len(t) <= 8 and not re.search(r"[A-Za-z]", t):
        return True
    if "\ufffd" in t:
        return True
    if re.fullmatch(r"[A-Z]\.?", t):
        return True
    if re.fullmatch(r"\d+\s*[•·]\s*", t):
        return True
    return False


def extract_page_lines(page: fitz.Page) -> Tuple[List[LineBlock], List[VisualBlock]]:
    data = page.get_text("dict")
    lines_out: List[LineBlock] = []
    visuals: List[VisualBlock] = []

    for blk in data.get("blocks", []):
        btype = blk.get("type", 0)
        if btype == 0:
            for ln in blk.get("lines", []):
                spans: List[SpanRun] = []
                pieces: List[str] = []
                max_font = 0.0
                font_sum = 0.0
                font_count = 0
                bold = False
                italic = False

                for sp in ln.get("spans", []):
                    txt = normalize_text(str(sp.get("text", "")))
                    if not txt:
                        continue
                    size = float(sp.get("size", 0.0))
                    font_name = str(sp.get("font", "")).lower()
                    sp_bold = "bold" in font_name
                    sp_italic = ("italic" in font_name) or ("oblique" in font_name)
                    spans.append(SpanRun(text=txt, size=size, bold=sp_bold, italic=sp_italic))
                    pieces.append(txt)
                    max_font = max(max_font, size)
                    font_sum += size
                    font_count += 1
                    bold = bold or sp_bold
                    italic = italic or sp_italic

                if not spans:
                    continue
                rect = fitz.Rect(ln["bbox"])
                text = normalize_text(" ".join(pieces))
                avg_font = (font_sum / font_count) if font_count else max_font
                lines_out.append(
                    LineBlock(
                        text=text,
                        rect=rect,
                        spans=spans,
                        max_font=max_font,
                        avg_font=avg_font,
                        is_bold=bold,
                        is_italic=italic,
                    )
                )
        elif btype == 1:
            visuals.append(VisualBlock(rect=fitz.Rect(blk["bbox"]), source="image"))

    for d in page.get_drawings():
        r = d.get("rect")
        if r is None:
            continue
        visuals.append(VisualBlock(rect=fitz.Rect(r), source="drawing"))

    return lines_out, visuals


def detect_two_column(lines: Sequence[LineBlock], page_w: float) -> bool:
    narrow = [ln for ln in lines if ln.rect.width < page_w * 0.58]
    left = [ln for ln in narrow if ln.rect.x0 < page_w * 0.45]
    right = [ln for ln in narrow if ln.rect.x0 >= page_w * 0.55]
    return len(left) >= 6 and len(right) >= 6


def order_lines(lines: List[LineBlock], page_w: float) -> List[LineBlock]:
    if not lines:
        return []
    if detect_two_column(lines, page_w):
        left = sorted([ln for ln in lines if ln.rect.x0 < page_w * 0.5], key=lambda x: (x.rect.y0, x.rect.x0))
        right = sorted([ln for ln in lines if ln.rect.x0 >= page_w * 0.5], key=lambda x: (x.rect.y0, x.rect.x0))
        return left + right
    return sorted(lines, key=lambda x: (x.rect.y0, x.rect.x0))


def classify_line(line: LineBlock, cfg: Dict[str, object], page_w: float, page_h: float) -> str:
    text = line.text
    if re.match(r"^\s*CHAPTER\s+\d+\s*$", text, re.I) and line.rect.y0 < page_h * 0.22:
        return "h1"
    if CAPTION_RE.search(text):
        if (
            (line.is_bold or len(text) <= 120)
            and not re.search(r"\bcompares?\b|\bdiscusses?\b|\bexplains?\b|\bdescribes?\b", text, re.I)
        ):
            return "caption"
        # In-text mentions like "Figure 1.1 compares..." should stay body text.
        return "body"

    heading_min = float(get_opt(cfg, "thresholds.heading_min_font_size"))
    if SECTION_RE.match(text):
        return "h2"
    if (
        (line.max_font >= heading_min or line.is_bold)
        and len(text) < 120
        and text.upper() == text
        and re.search(r"[A-Z]", text)
        and not CAPTION_RE.search(text)
    ):
        return "h2"

    eq_density = float(get_opt(cfg, "thresholds.equation_symbol_density"))
    eq_max_chars = int(get_opt(cfg, "thresholds.equation_max_chars"))
    center = abs(((line.rect.x0 + line.rect.x1) * 0.5) - (page_w * 0.5)) < page_w * 0.18
    math_hits = len(re.findall(r"[=<>∑∫√]|(?:\bmax\b)|(?:\bmin\b)", text, re.I))
    token_count = len(text.split())
    long_alpha_words = len(re.findall(r"\b[A-Za-z]{4,}\b", text))
    sentence_like = long_alpha_words >= 4 and bool(
        re.search(
            r"\b(?:the|and|for|with|that|this|to|of|an|in|on|at|from|by|is|are|was|were|will|would|be)\b",
            text,
            re.I,
        )
    )
    if (
        len(text) <= min(eq_max_chars, 105)
        and center
        and token_count <= 24
        and math_hits >= 1
        and symbol_density(text) >= eq_density
        and not sentence_like
    ):
        return "equation"

    return "body"


def combine_rects(rects: Iterable[fitz.Rect]) -> Optional[fitz.Rect]:
    rects = list(rects)
    if not rects:
        return None
    out = fitz.Rect(rects[0])
    for r in rects[1:]:
        out.include_rect(r)
    return out


def clamp_rect(rect: fitz.Rect, page_rect: fitz.Rect) -> fitz.Rect:
    return fitz.Rect(
        max(page_rect.x0, rect.x0),
        max(page_rect.y0, rect.y0),
        min(page_rect.x1, rect.x1),
        min(page_rect.y1, rect.y1),
    )


def intersects_ratio(a: fitz.Rect, b: fitz.Rect) -> float:
    inter = a & b
    if inter.is_empty:
        return 0.0
    denom = max(a.get_area(), 1.0)
    return inter.get_area() / denom


def find_figure_region(
    caption: LineBlock,
    lines: Sequence[LineBlock],
    visuals: Sequence[VisualBlock],
    page_rect: fitz.Rect,
    cfg: Dict[str, object],
) -> Tuple[Optional[fitz.Rect], bool]:
    x_margin = page_rect.width * float(get_opt(cfg, "thresholds.figure_x_margin_ratio"))
    y_extension = page_rect.height * float(get_opt(cfg, "thresholds.figure_y_extension_ratio"))
    min_visual_area = page_rect.get_area() * float(get_opt(cfg, "thresholds.min_visual_area_ratio"))

    # Preferred heuristic from user feedback:
    # figure/table regions are wrapped by two long horizontal gray strips.
    strip_like = [
        v.rect
        for v in visuals
        if v.source == "drawing"
        and v.rect.width >= page_rect.width * 0.68
        and v.rect.height <= page_rect.height * 0.02
    ]
    top_strip = None
    bottom_strip = None
    for r in sorted(strip_like, key=lambda rr: rr.y0):
        if r.y1 <= caption.rect.y0 + page_rect.height * 0.05:
            top_strip = r
        elif r.y0 > caption.rect.y1 and bottom_strip is None:
            bottom_strip = r
            break
    if top_strip is not None and bottom_strip is not None and bottom_strip.y0 > top_strip.y1:
        strip_x0 = max(page_rect.x0, min(top_strip.x0, bottom_strip.x0) - 1)
        strip_x1 = min(page_rect.x1, max(top_strip.x1, bottom_strip.x1) + 1)
        bounded = fitz.Rect(
            strip_x0,
            max(top_strip.y1 + 1, page_rect.y0),
            strip_x1,
            min(bottom_strip.y0 - 1, page_rect.y1),
        )
        if bounded.get_area() > page_rect.get_area() * 0.06 and bounded.get_area() < page_rect.get_area() * 0.58:
            return clamp_rect(bounded, page_rect), True

    # Find next strong boundary (new section / next caption / clear body start far below).
    stop_y = page_rect.y1 - page_rect.height * 0.03
    for ln in lines:
        if ln.rect.y0 <= caption.rect.y1 + 2:
            continue
        t = ln.text
        if SECTION_RE.match(t) or CAPTION_RE.search(t):
            stop_y = min(stop_y, ln.rect.y0 - 4)
            break
        # Stop when likely entering normal body paragraph text.
        if (
            len(t) > 72
            and ln.avg_font >= 9.5
            and symbol_density(t) < 0.03
            and not re.match(r"^\s*(?:Figure|Table)\b", t, re.I)
        ):
            stop_y = min(stop_y, ln.rect.y0 - 3)
            break

    # Candidate visual elements in the window beneath caption.
    v_candidates = [
        v.rect
        for v in visuals
        if v.rect.get_area() >= min_visual_area
        and v.rect.y1 >= caption.rect.y1
        and v.rect.y0 <= caption.rect.y1 + y_extension
    ]

    text_label_candidates = [
        ln.rect
        for ln in lines
        if ln.rect.y0 >= caption.rect.y1
        and ln.rect.y0 < stop_y
        and (len(ln.text) <= 32 or ln.avg_font <= 9.8)
        and not CAPTION_RE.search(ln.text)
    ]

    base_rect = combine_rects(v_candidates + text_label_candidates)
    if base_rect is None:
        # Bounded fallback for vector-heavy figures where drawing extraction is weak.
        fallback = fitz.Rect(
            page_rect.x0 + x_margin,
            max(caption.rect.y0 - 2, page_rect.y0),
            page_rect.x1 - x_margin,
            min(stop_y, caption.rect.y1 + page_rect.height * 0.45),
        )
        if fallback.height >= page_rect.height * 0.12 and fallback.width >= page_rect.width * 0.35:
            return clamp_rect(fallback, page_rect), False
        return None, False

    expanded = fitz.Rect(
        page_rect.x0 + x_margin,
        max(caption.rect.y0 - 2, base_rect.y0 - 6),
        page_rect.x1 - x_margin,
        min(stop_y, base_rect.y1 + 8),
    )

    if expanded.width < page_rect.width * 0.35 or expanded.height < page_rect.height * 0.08:
        return None, False

    # Never allow a near-full-page crop for "figure" regions.
    if expanded.get_area() > page_rect.get_area() * 0.58:
        expanded = fitz.Rect(
            expanded.x0,
            expanded.y0,
            expanded.x1,
            min(expanded.y1, caption.rect.y1 + page_rect.height * 0.48),
        )
        if expanded.get_area() > page_rect.get_area() * 0.58:
            return None, False

    return clamp_rect(expanded, page_rect), False


def line_to_inline_html(line: LineBlock, body_font: float) -> str:
    parts: List[str] = []
    for sp in line.spans:
        txt = html.escape(sp.text)
        if not txt:
            continue

        classes: List[str] = []
        rel = sp.size / body_font if body_font > 0 else 1.0
        if rel <= 0.92:
            classes.append("fs-small")
        elif rel >= 1.12:
            classes.append("fs-large")

        wrapped = txt
        if sp.italic:
            wrapped = f"<em>{wrapped}</em>"
        if sp.bold:
            wrapped = f"<strong>{wrapped}</strong>"
        if classes:
            wrapped = f'<span class="{" ".join(classes)}">{wrapped}</span>'
        parts.append(wrapped)

    joined = " ".join(parts)
    joined = re.sub(r"\s+", " ", joined).strip()
    return joined


def cleanup_equation_text(s: str) -> str:
    s = re.sub(r"\s{2,}", " ", s).strip()
    # Common PDF extraction artifacts in this chapter.
    s = re.sub(r"\bS\s*T\b", "S_T", s)
    s = s.replace("S_T 7 K", "S_T > K")
    s = s.replace("S_T ... K", "S_T <= K")
    s = re.sub(r"\bmax\s*1\s*([^,]+),\s*0\s*2\b", r"max(\1, 0)", s)
    s = re.sub(r"\bmin\s*1\s*([^,]+),\s*0\s*2\b", r"min(\1, 0)", s)
    s = re.sub(r"\bmax1\s*([^,]+),\s*0\s*2\b", r"max(\1, 0)", s)
    s = re.sub(r"\bmin1\s*([^,]+),\s*0\s*2\b", r"min(\1, 0)", s)
    s = re.sub(r",\s*0\s*2\b", ", 0)", s)
    s = re.sub(r"\s*-\s*max\(", " -max(", s)
    s = re.sub(r"\s*=\s*min\(", " = min(", s)
    s = normalize_inline_math_ocr(s)
    return s


def normalize_inline_math_ocr(s: str) -> str:
    # Common OCR artifact in this PDF: "(= ... )" appears as "1= ... 2".
    def repl(m: re.Match[str]) -> str:
        a, b, c = m.group(1), m.group(2), m.group(3)
        syms = [x[0] for x in (a, b, c) if x and x[0] in {"$", "£"}]
        sym = syms[0] if syms else "$"
        def ensure_cur(x: str) -> str:
            return x if x and x[0] in {"$", "£"} else f"{sym}{x}"
        return f"{ensure_cur(a)} (= {ensure_cur(b)} - {ensure_cur(c)})"

    s = re.sub(
        r"\+\s*([£$]?\d[\d,]*)\s*1\s*=\s*\+\s*([£$]?\d[\d,]*)\s*-\s*\+\s*([£$]?\d[\d,]*)\s*2",
        repl,
        s,
    )
    s = re.sub(
        r"([£$]?\d[\d,]*)\s*1\s*=\s*([£$]?\d[\d,]*)\s*-\s*([£$]?\d[\d,]*)\s*2",
        repl,
        s,
    )
    return s


def trim_image_whitespace(path: Path, threshold: int = 8, padding: int = 4) -> None:
    if Image is None:
        return
    img = Image.open(path).convert("RGB")
    w, h = img.size
    # Sample corners to estimate background color.
    corners = [
        img.getpixel((0, 0)),
        img.getpixel((w - 1, 0)),
        img.getpixel((0, h - 1)),
        img.getpixel((w - 1, h - 1)),
    ]
    bg = tuple(int(sum(c[i] for c in corners) / 4) for i in range(3))
    px = img.load()

    minx, miny = w, h
    maxx, maxy = -1, -1
    for y in range(h):
        for x in range(w):
            r, g, b = px[x, y]
            if abs(r - bg[0]) + abs(g - bg[1]) + abs(b - bg[2]) > threshold:
                if x < minx:
                    minx = x
                if y < miny:
                    miny = y
                if x > maxx:
                    maxx = x
                if y > maxy:
                    maxy = y

    if maxx <= minx or maxy <= miny:
        return
    crop = (
        max(0, minx - padding),
        max(0, miny - padding),
        min(w, maxx + 1 + padding),
        min(h, maxy + 1 + padding),
    )
    if crop == (0, 0, w, h):
        return
    # Ignore tiny changes to avoid churn.
    if (crop[2] - crop[0]) >= int(w * 0.98) and (crop[3] - crop[1]) >= int(h * 0.98):
        return
    img.crop(crop).save(path)


def crop_below_bottom_gray_bar(path: Path) -> None:
    if Image is None:
        return
    img = Image.open(path).convert("RGB")
    w, h = img.size
    px = img.load()

    def is_grayish(r: int, g: int, b: int) -> bool:
        avg = (r + g + b) // 3
        return 125 <= avg <= 235 and abs(r - g) <= 10 and abs(g - b) <= 10

    def gray_run_ratio(y: int) -> Tuple[float, float]:
        gray = 0
        best_run = 0
        run = 0
        for x in range(w):
            if is_grayish(*px[x, y]):
                gray += 1
                run += 1
                if run > best_run:
                    best_run = run
            else:
                run = 0
        return gray / max(w, 1), best_run / max(w, 1)

    # Find horizontal bands with strong gray bars (works even with side margins).
    candidates: List[Tuple[int, float]] = []
    for y in range(int(h * 0.25), h):
        ratio, run_ratio = gray_run_ratio(y)
        if ratio >= 0.58 and run_ratio >= 0.55:
            candidates.append((y, run_ratio))

    if not candidates:
        return

    # Merge contiguous rows to bands.
    bands: List[Tuple[int, int, float]] = []
    start = candidates[0][0]
    prev = start
    ratios = [candidates[0][1]]
    for y, r in candidates[1:]:
        if y == prev + 1:
            prev = y
            ratios.append(r)
        else:
            bands.append((start, prev, sum(ratios) / len(ratios)))
            start, prev, ratios = y, y, [r]
    bands.append((start, prev, sum(ratios) / len(ratios)))

    # Pick the lowest substantial band (the panel bottom bar).
    target = None
    for b in bands:
        y0, y1, _ = b
        if (y1 - y0 + 1) >= 2 and y0 >= int(h * 0.45):
            target = b
    if target is None:
        return

    y0, y1, _ = target
    cut_y = min(h, y1 + 3)
    if cut_y < int(h * 0.55):
        return
    if cut_y >= h:
        return
    img.crop((0, 0, w, cut_y)).save(path)


def gray_bar_trailing_pixels(path: Path) -> int:
    if Image is None:
        return 0
    img = Image.open(path).convert("RGB")
    w, h = img.size
    px = img.load()

    def is_grayish(r: int, g: int, b: int) -> bool:
        avg = (r + g + b) // 3
        return 125 <= avg <= 235 and abs(r - g) <= 10 and abs(g - b) <= 10

    def gray_run_ratio(y: int) -> Tuple[float, float]:
        gray = 0
        best_run = 0
        run = 0
        for x in range(w):
            if is_grayish(*px[x, y]):
                gray += 1
                run += 1
                if run > best_run:
                    best_run = run
            else:
                run = 0
        return gray / max(w, 1), best_run / max(w, 1)

    rows = []
    for y in range(int(h * 0.25), h):
        ratio, run_ratio = gray_run_ratio(y)
        if ratio >= 0.58 and run_ratio >= 0.55:
            rows.append(y)
    if not rows:
        return 0
    near_bottom = [y for y in rows if y >= int(h * 0.70)]
    if not near_bottom:
        return 0
    bottom = max(near_bottom)
    bg = img.getpixel((0, 0))
    trail = 0
    for y in range(bottom + 1, h):
        for x in range(w):
            r, g, b = px[x, y]
            if abs(r - bg[0]) + abs(g - bg[1]) + abs(b - bg[2]) > 10:
                trail += 1
    return trail


def equation_to_html(raw: str) -> str:
    eq = cleanup_equation_text(raw)
    eq = eq.replace("<=", "≤").replace(">=", "≥").replace("!=", "≠")
    esc = html.escape(eq)
    esc = re.sub(r"([A-Za-z])_([A-Za-z0-9]+)", r"<em>\1</em><sub>\2</sub>", esc)
    esc = re.sub(r"([A-Za-z0-9\)])\^([A-Za-z0-9+\-]+)", r"\1<sup>\2</sup>", esc)
    return esc


def body_font_size(lines: Sequence[LineBlock]) -> float:
    sizes = [ln.avg_font for ln in lines if 8.0 <= ln.avg_font <= 14.5]
    if not sizes:
        return 11.0
    return statistics.median(sizes)


def compose_content_blocks_for_page(
    lines: List[LineBlock],
    visuals: List[VisualBlock],
    page: fitz.Page,
    page_num: int,
    chapter_id: str,
    cfg: Dict[str, object],
    img_dir: Path,
    figure_start_idx: int,
) -> Tuple[List[ContentBlock], int, Dict[str, int]]:
    page_rect = page.rect
    lines = [ln for ln in lines if not is_header_footer(ln, page_rect.height)]
    lines = order_lines(lines, page_rect.width)

    body_font = body_font_size(lines)
    body_lines = [ln for ln in lines if classify_line(ln, cfg, page_rect.width, page_rect.height) == "body"]
    gaps = []
    for i in range(1, len(body_lines)):
        g = body_lines[i].rect.y0 - body_lines[i - 1].rect.y1
        if g > 0:
            gaps.append(g)
    median_gap = statistics.median(gaps) if gaps else 2.4
    para_gap_threshold = median_gap * float(get_opt(cfg, "thresholds.paragraph_gap_multiplier"))

    excluded_lines: List[fitz.Rect] = []
    blocks: List[ContentBlock] = []
    figure_counter = figure_start_idx
    stats = {
        "captions": 0,
        "figures": 0,
        "figure_suppressed_lines": 0,
        "long_paragraphs": 0,
        "style_runs_italic": 0,
        "style_runs_bold": 0,
    }

    pending_para: List[LineBlock] = []

    def flush_para() -> None:
        nonlocal pending_para
        if not pending_para:
            return
        txt = " ".join(ln.text for ln in pending_para).strip()
        if len(txt) > 900:
            stats["long_paragraphs"] += 1
        r = combine_rects([ln.rect for ln in pending_para]) or fitz.Rect(page_rect)
        blocks.append(ContentBlock(kind="p", rect=r, lines=list(pending_para), text=txt, page_num=page_num))
        pending_para = []

    def is_excluded(line: LineBlock) -> bool:
        return any(intersects_ratio(line.rect, ex) >= 0.35 for ex in excluded_lines)

    dpi = int(get_opt(cfg, "dpi"))
    zoom = dpi / 72.0

    i = 0
    while i < len(lines):
        ln = lines[i]
        if is_excluded(ln):
            i += 1
            continue

        kind = classify_line(ln, cfg, page_rect.width, page_rect.height)

        if kind in {"caption", "h1", "h2", "equation"}:
            flush_para()

        if kind == "caption":
            stats["captions"] += 1
            # Extend caption with wrapped continuation lines near the top of the figure panel.
            caption_lines = [ln]
            j = i + 1
            while j < len(lines):
                cand = lines[j]
                if CAPTION_RE.search(cand.text) or SECTION_RE.match(cand.text):
                    break
                if cand.rect.y0 - caption_lines[-1].rect.y1 > page_rect.height * 0.02:
                    break
                if cand.rect.y0 > ln.rect.y0 + page_rect.height * 0.12:
                    break
                if len(cand.text) > 140:
                    break
                # Keep nearby lines that align with caption block.
                if abs(cand.rect.x0 - ln.rect.x0) <= page_rect.width * 0.08:
                    caption_lines.append(cand)
                    j += 1
                    continue
                break

            caption_rect = combine_rects([c.rect for c in caption_lines]) or ln.rect
            caption_text = " ".join(c.text for c in caption_lines).strip()
            # Drop common axis-label leakage from plot area.
            caption_text = re.sub(r"\s+Profit\s*\(\$\)\s*$", "", caption_text, flags=re.I)
            caption_text = re.sub(r"\s+Terminal\s+stock\s+price\s*\(\$\)\s*$", "", caption_text, flags=re.I)
            next_caption_y = page_rect.y1
            for look in lines[j:]:
                if CAPTION_RE.search(look.text):
                    next_caption_y = look.rect.y0 - 3
                    break
            region, strip_wrapped = find_figure_region(ln, lines[i + 1 :], visuals, page_rect, cfg)
            if region is not None:
                # Keep caption as text; only figure body becomes image.
                body_region = fitz.Rect(region)
                caption_above = caption_rect.y0 <= (region.y0 + region.height * 0.35)
                if strip_wrapped and caption_above:
                    body_region.y0 = max(body_region.y0, caption_rect.y1 + 2)
                    body_region.y1 = min(body_region.y1, next_caption_y)
                elif caption_above:
                    body_region.y1 = min(body_region.y1, next_caption_y)
                else:
                    body_region.y1 = min(body_region.y1, caption_rect.y0 - 2, next_caption_y)
                if body_region.height < page_rect.height * 0.06:
                    body_region = region
                # Tighten body crop to plotted content so image can fill EPUB width better.
                content_rects: List[fitz.Rect] = []
                for v in visuals:
                    if intersects_ratio(v.rect, body_region) > 0.01:
                        content_rects.append(v.rect)
                for cand in lines:
                    if cand.rect.y0 < body_region.y0 or cand.rect.y1 > body_region.y1:
                        continue
                    if len(cand.text) <= 40 or symbol_density(cand.text) > 0.03:
                        content_rects.append(cand.rect)
                tight = combine_rects(content_rects)
                if tight is not None and tight.width > page_rect.width * 0.30:
                    # Avoid top clipping on plot titles/labels: keep original top for strip-wrapped figures.
                    top_y = body_region.y0 if strip_wrapped else max(body_region.y0, tight.y0 - 4)
                    body_region = fitz.Rect(
                        body_region.x0,
                        top_y,
                        body_region.x1,
                        min(body_region.y1, tight.y1 + 6),
                    )
                figure_counter += 1
                fig_id = f"fig-{chapter_id}-{figure_counter}"
                pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), clip=body_region, alpha=False)
                img_name = f"p{page_num}_fig{figure_counter}.png"
                img_path = img_dir / img_name
                pix.save(img_path.as_posix())
                crop_below_bottom_gray_bar(img_path)
                trim_image_whitespace(img_path)

                excluded_now = 0
                for other in lines:
                    if other.rect.y0 < ln.rect.y1:
                        continue
                    if intersects_ratio(other.rect, region) >= 0.50:
                        excluded_lines.append(other.rect)
                        excluded_now += 1
                stats["figure_suppressed_lines"] += excluded_now

                blocks.append(
                    ContentBlock(
                        kind="figure",
                        rect=body_region,
                        text=caption_text,
                        page_num=page_num,
                        fig_id=fig_id,
                        img_name=f"images/{img_name}",
                        figure_has_inline_caption=False,
                    )
                )
                stats["figures"] += 1
                for c in caption_lines:
                    excluded_lines.append(c.rect)
                # Advance past consumed caption continuation lines.
                i = j - 1
            else:
                # Bounded fallback: for real caption lines where region detection is weak,
                # crop a limited panel below the caption (never full-page).
                fallback_stop = page_rect.y1 - page_rect.height * 0.04
                for look in lines[j:]:
                    if CAPTION_RE.search(look.text) or SECTION_RE.match(look.text):
                        fallback_stop = min(fallback_stop, look.rect.y0 - 3)
                        break
                start_y = caption_rect.y1 + 2
                force_region = fitz.Rect(
                    page_rect.x0 + page_rect.width * 0.02,
                    start_y,
                    page_rect.x1 - page_rect.width * 0.02,
                    min(fallback_stop, start_y + page_rect.height * 0.48),
                )
                if force_region.height >= page_rect.height * 0.10 and force_region.width >= page_rect.width * 0.40:
                    figure_counter += 1
                    fig_id = f"fig-{chapter_id}-{figure_counter}"
                    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), clip=force_region, alpha=False)
                    img_name = f"p{page_num}_fig{figure_counter}.png"
                    img_path = img_dir / img_name
                    pix.save(img_path.as_posix())
                    crop_below_bottom_gray_bar(img_path)
                    trim_image_whitespace(img_path)
                    blocks.append(
                        ContentBlock(
                            kind="figure",
                            rect=force_region,
                            text=caption_text,
                            page_num=page_num,
                            fig_id=fig_id,
                            img_name=f"images/{img_name}",
                            figure_has_inline_caption=False,
                        )
                    )
                    stats["figures"] += 1
                    for c in caption_lines:
                        excluded_lines.append(c.rect)
                    for other in lines:
                        if other.rect.y0 < caption_rect.y1:
                            continue
                        if intersects_ratio(other.rect, force_region) >= 0.50:
                            excluded_lines.append(other.rect)
                            stats["figure_suppressed_lines"] += 1
                else:
                    # If no robust figure region found, keep caption as normal paragraph text.
                    blocks.append(
                        ContentBlock(kind="p", rect=caption_rect, lines=caption_lines, text=caption_text, page_num=page_num)
                    )
                i = j - 1

        elif kind == "h1":
            blocks.append(ContentBlock(kind="h1", rect=ln.rect, lines=[ln], text=ln.text, page_num=page_num))

        elif kind == "h2":
            blocks.append(ContentBlock(kind="h2", rect=ln.rect, lines=[ln], text=ln.text, page_num=page_num))

        elif kind == "equation":
            blocks.append(ContentBlock(kind="equation", rect=ln.rect, lines=[ln], text=ln.text, page_num=page_num))

        else:
            if not pending_para:
                pending_para.append(ln)
            else:
                prev = pending_para[-1]
                y_gap = ln.rect.y0 - prev.rect.y1
                x_shift = abs(ln.rect.x0 - pending_para[0].rect.x0)
                prev_end = prev.text.strip()
                starts_numbered = bool(re.match(r"^\d+\.\d+", ln.text.strip()))
                new_para = False

                if y_gap > para_gap_threshold:
                    new_para = True
                if x_shift > page_rect.width * 0.08 and SENTENCE_END_RE.search(prev_end):
                    new_para = True
                if (
                    prev.rect.x1 < page_rect.x1 - page_rect.width * 0.12
                    and abs(ln.rect.x0 - pending_para[0].rect.x0) < page_rect.width * 0.03
                    and SENTENCE_END_RE.search(prev_end)
                ):
                    new_para = True
                if y_gap > median_gap * 0.6 and SENTENCE_END_RE.search(prev_end):
                    if re.match(r"^[A-Z(]", ln.text.strip()):
                        new_para = True
                if starts_numbered:
                    new_para = True

                if new_para:
                    flush_para()
                    pending_para.append(ln)
                else:
                    pending_para.append(ln)

        for sp in ln.spans:
            if sp.italic:
                stats["style_runs_italic"] += 1
            if sp.bold:
                stats["style_runs_bold"] += 1

        i += 1

    flush_para()

    # Promote equation-like standalone paragraph lines to equation blocks.
    for idx, b in enumerate(blocks):
        if b.kind != "p" or len(b.lines) != 1:
            continue
        t = b.lines[0].text.strip()
        tok = len(t.split())
        alpha_words = len(re.findall(r"\b[A-Za-z]{4,}\b", t))
        if (
            2 <= tok <= 12
            and alpha_words <= 2
            and symbol_density(t) >= 0.07
            and re.search(r"[=<>+\-*/()]", t)
        ):
            blocks[idx].kind = "equation"

    # Convert line-based blocks to HTML-ready text where needed.
    for b in blocks:
        if b.kind in {"p", "equation", "h1", "h2"} and b.lines:
            if b.kind == "p":
                # Keep paragraph reflowable as true paragraphs instead of line-by-line breaks.
                b.text = " ".join(line_to_inline_html(ln, body_font) for ln in b.lines)
                # Dehyphenate PDF line-break artifacts: "Con- sider" -> "Consider".
                b.text = re.sub(r"([A-Za-z])-\s+([a-z])", r"\1\2", b.text)
                b.text = re.sub(r"\s{2,}", " ", b.text).strip()
                b.text = normalize_inline_math_ocr(b.text)
                # Normalize common math inline artifacts.
                b.text = b.text.replace(
                    "<em>S</em> <span class=\"fs-small\"><em>T</em></span>",
                    "<em>S</em><sub>T</sub>",
                )
                b.text = re.sub(r"(<em>S</em><sub>T</sub>)\s*7\s*(<em>K</em>)", r"\1 &gt; \2", b.text)
                b.text = re.sub(r"(<em>S</em><sub>T</sub>)\s*…\s*(<em>K</em>)", r"\1 &lt;= \2", b.text)
            elif b.kind == "equation":
                raw = " ".join(ln.text for ln in b.lines)
                b.text = equation_to_html(raw)
            else:
                b.text = " ".join(line_to_inline_html(ln, body_font) for ln in b.lines)

        if b.kind == "figure" and b.text:
            b.text = re.sub(r"(Figure\s+\d+\.\d)([A-Za-z])", r"\1 \2", b.text)

    return blocks, figure_counter, stats


def build_css() -> str:
    return """
body {
  margin: 0;
  padding: 0 0.75rem;
  line-height: 1.56;
  font-family: Georgia, "Times New Roman", serif;
  color: #1f1f1f;
  text-align: justify !important;
  text-justify: inter-word;
  -webkit-hyphens: auto;
}

h1 {
  font-size: 1.52rem;
  line-height: 1.25;
  margin: 1.2rem 0 0.8rem;
}

h2 {
  font-size: 1.08rem;
  line-height: 1.34;
  margin: 1.0rem 0 0.55rem;
  font-weight: 600;
}

p {
  margin: 0.42rem 0;
  text-indent: 0;
  text-align: justify !important;
  text-justify: inter-word;
  hyphens: auto;
  -webkit-hyphens: auto;
}

.equation {
  font-family: "Times New Roman", Georgia, serif;
  display: block;
  width: 100%;
  text-align: center !important;
  white-space: normal;
  margin: 0.9rem auto;
  font-size: 1.02rem;
  line-height: 1.5;
  max-width: 100%;
}

.equation .eq-inline {
  display: inline-block;
  margin: 0 auto;
  text-align: center !important;
}

.fs-small {
  font-size: 0.9em;
}

.fs-large {
  font-size: 1.08em;
}

.fig-block {
  display: block;
  margin: 0.9rem 0;
  width: 100% !important;
  max-width: 100% !important;
  overflow: visible !important;
  -webkit-column-span: all;
  column-span: all;
  break-inside: avoid;
}

.fig-cap {
  display: block;
  width: 100% !important;
  margin: 0 0 0.30rem 0;
  font-size: 0.86rem;
  line-height: 1.34;
  color: #3f3f3f;
  text-align: justify !important;
  text-justify: inter-word;
  white-space: normal !important;
  overflow: visible !important;
  overflow-wrap: break-word;
}

.fig-img {
  display: block;
  width: 100% !important;
  max-width: 100% !important;
  height: auto !important;
  margin: 0;
}

.toc-brief {
  margin: 0.8rem 0 0.4rem;
  padding-left: 1.2rem;
}

.toc-brief li {
  margin: 0.26rem 0;
  text-align: left !important;
  line-height: 1.35;
}

.toc-brief a {
  text-decoration: none;
}

p.toc-line {
  text-align: left !important;
  text-justify: auto !important;
  hyphens: none !important;
  -webkit-hyphens: none !important;
  margin: 0.18rem 0;
}

.toc-main {
  column-count: 2;
  column-gap: 2rem;
  margin: 0.8rem 0 0.5rem;
}

.toc-main .toc-main-chapter {
  break-inside: avoid;
  margin: 0.55rem 0 0.18rem;
  font-weight: 700;
  text-align: left !important;
}

.toc-main .toc-main-line {
  break-inside: avoid;
  margin: 0.10rem 0;
  text-align: left !important;
}

.toc-main .toc-main-sec {
  margin-left: 1.1rem;
}

.toc-main a {
  text-decoration: none;
}

.list-page {
  margin: 0.8rem 0 0.5rem;
}

.list-page .list-item {
  margin: 0.14rem 0;
  text-align: left !important;
}

.list-page a {
  text-decoration: none;
}
""".strip()


def chapter_to_xhtml(
    chapter_title: str,
    blocks: List[ContentBlock],
    current_page: int,
    page_file_map: Dict[int, str],
    figure_table_targets: Dict[str, Tuple[int, str]],
    chapter_targets: Dict[str, Tuple[int, str]],
    toc_override: Optional[List[Tuple[str, str]]] = None,
    toc_main_override: Optional[List[Tuple[int, str, str]]] = None,
    business_snapshots_override: Optional[List[Tuple[str, str]]] = None,
    technical_notes_override: Optional[List[str]] = None,
) -> str:
    parts: List[str] = []
    parts.append("<html>")
    parts.append("<head>")
    parts.append('<meta charset="utf-8" />')
    parts.append(f"<title>{html.escape(chapter_title)}</title>")
    parts.append('<link rel="stylesheet" type="text/css" href="style.css" />')
    parts.append("</head><body>")
    # Render clean text TOC page from official chapter list when available.
    if toc_override:
        parts.append("<h1>CONTENTS IN BRIEF</h1>")
        parts.append('<ol class="toc-brief">')
        for label, href in toc_override:
            parts.append(
                f'<li><a href="{escape_attr(href)}">{html.escape(label)}</a></li>'
            )
        parts.append("</ol>")
        parts.append("</body></html>")
        return "\n".join(parts)

    if toc_main_override is not None:
        parts.append("<h1>CONTENTS</h1>")
        parts.append('<div class="toc-main">')
        for lvl, label, href in toc_main_override:
            cls = "toc-main-chapter" if lvl == 1 else "toc-main-line toc-main-sec"
            parts.append(
                f'<div class="{cls}"><a href="{escape_attr(href)}">{html.escape(label)}</a></div>'
            )
        parts.append("</div>")
        parts.append("</body></html>")
        return "\n".join(parts)

    if business_snapshots_override is not None:
        parts.append("<h1>BUSINESS SNAPSHOTS</h1>")
        parts.append('<div class="list-page">')
        for label, href in business_snapshots_override:
            parts.append(
                f'<div class="list-item"><a href="{escape_attr(href)}">{html.escape(label)}</a></div>'
            )
        parts.append("</div>")
        parts.append("</body></html>")
        return "\n".join(parts)

    if technical_notes_override is not None:
        parts.append("<h1>TECHNICAL NOTES</h1>")
        parts.append('<div class="list-page">')
        for line in technical_notes_override:
            parts.append(f'<div class="list-item">{html.escape(line)}</div>')
        parts.append("</div>")
        parts.append("</body></html>")
        return "\n".join(parts)

    # Do not inject repeated static h1 on every page; headings come from detected blocks.

    for b in blocks:
        if b.kind == "h1":
            htxt = linkify_html_text(
                b.text,
                current_page=current_page,
                page_file_map=page_file_map,
                figure_table_targets=figure_table_targets,
                chapter_targets=chapter_targets,
            )
            hid = f' id="{escape_attr(b.anchor_id)}"' if b.anchor_id else ""
            parts.append(f"<h1{hid}>{htxt}</h1>")
        elif b.kind == "h2":
            htxt = linkify_html_text(
                b.text,
                current_page=current_page,
                page_file_map=page_file_map,
                figure_table_targets=figure_table_targets,
                chapter_targets=chapter_targets,
            )
            hid = f' id="{escape_attr(b.anchor_id)}"' if b.anchor_id else ""
            parts.append(f"<h2{hid}>{htxt}</h2>")
        elif b.kind == "equation":
            etxt = linkify_html_text(
                b.text,
                current_page=current_page,
                page_file_map=page_file_map,
                figure_table_targets=figure_table_targets,
                chapter_targets=chapter_targets,
            )
            parts.append(
                '<div class="equation" style="text-align:center !important; width:100%;">'
                f'<span class="eq-inline">{etxt}</span></div>'
            )
        elif b.kind == "figure":
            alt = escape_attr(b.text or "Figure")
            caption = linkify_html_text(
                html.escape(b.text),
                current_page=current_page,
                page_file_map=page_file_map,
                figure_table_targets=figure_table_targets,
                chapter_targets=chapter_targets,
            )
            img = escape_attr(b.img_name or "")
            fig_id = escape_attr(b.fig_id or "")
            if b.figure_has_inline_caption:
                parts.append(
                    f'<div id="{fig_id}" class="fig-block">'
                    f'<img class="fig-img" src="{img}" alt="{alt}" />'
                    "</div>"
                )
            else:
                parts.append(
                    f'<div id="{fig_id}" class="fig-block">'
                    f'<div class="fig-cap">{caption}</div>'
                    f'<img class="fig-img" src="{img}" alt="{alt}" />'
                    "</div>"
                )
        else:
            ptxt = linkify_html_text(
                b.text,
                current_page=current_page,
                page_file_map=page_file_map,
                figure_table_targets=figure_table_targets,
                chapter_targets=chapter_targets,
            )
            pclass = f' class="{escape_attr(b.css_class)}"' if b.css_class else ""
            parts.append(
                f'<p{pclass} style="text-align: justify; text-justify: inter-word; '
                '-webkit-hyphens: auto; hyphens: auto;">'
                f"{ptxt}</p>"
            )

    parts.append("</body></html>")
    return "\n".join(parts)


def build_reference_maps(blocks: List[ContentBlock]) -> Tuple[Dict[str, Tuple[int, str]], Dict[str, Tuple[int, str]]]:
    targets: Dict[str, Tuple[int, str]] = {}
    chapter_targets: Dict[str, Tuple[int, str]] = {}
    toc_pages = {
        b.page_num
        for b in blocks
        if b.kind in {"h1", "h2", "p"} and ("CONTENTS IN BRIEF" in html_to_plain_text(b.text).upper() or html_to_plain_text(b.text).upper() == "CONTENTS")
    }

    sec_counter = 0
    for b in blocks:
        if b.kind == "h1":
            if b.page_num in toc_pages:
                continue
            if not b.anchor_id:
                sec_counter += 1
                b.anchor_id = f"chap-{sec_counter}"
            txt = html_to_plain_text(b.text)
            m = re.search(r"\bCHAPTER\s+(\d+)\b", txt, re.I)
            if m:
                chapter_targets.setdefault(f"chapter {m.group(1)}", (b.page_num, b.anchor_id))
        if b.kind == "h2":
            if b.page_num in toc_pages:
                continue
            if not b.anchor_id:
                sec_counter += 1
                b.anchor_id = f"sec-{sec_counter}"

        if b.kind == "figure" and b.fig_id:
            lbl = html_to_plain_text(b.text)
            m = re.match(r"^\s*((?:Figure|Table)\s+\d+\.\d+)\b", lbl, re.I)
            if m:
                targets[m.group(1).lower()] = (b.page_num, b.fig_id)

    return targets, chapter_targets


def extract_contents_chapter_toc(doc: fitz.Document) -> List[Tuple[int, str]]:
    """
    Parse official chapter list from the book's CONTENTS pages.
    Returns [(chapter_number, chapter_title), ...].
    """
    out: List[Tuple[int, str]] = []
    seen = set()
    max_scan = min(len(doc), 40)
    for p in range(max_scan):
        txt = doc[p].get_text("text")
        if "CONTENTS" not in txt.upper():
            continue
        pending_chap: Optional[int] = None
        pending_title: str = ""
        for raw in txt.splitlines():
            line = raw.replace("\t", " ")
            line = re.sub(r"[\u0000-\u001F]", " ", line)
            line = re.sub(r"\s+", " ", line).strip()
            if not line:
                continue
            if re.fullmatch(r"(?:CONTENTS(?: IN BRIEF)?|Contents)", line, re.I):
                continue
            if line.isdigit():
                continue
            # Prefer chapter lines from "CONTENTS IN BRIEF", e.g. "10. Mechanics ... 227".
            m = re.match(r"^(\d{1,2})\.\s*(.+?)\s*\.{2,}\s*(\d+)\s*$", line)
            if not m:
                m = re.match(r"^(\d{1,2})\.\s*(.+?)\s+(\d+)\s*$", line)
            if not m:
                # Also support "Chapter 10. Mechanics ... 227" style.
                m2 = re.match(r"^Chapter\s+(\d{1,2})\.\s*(.+?)\s*\.{2,}\s*(\d+)\s*$", line, re.I)
                if not m2:
                    m2 = re.match(r"^Chapter\s+(\d{1,2})\.\s*(.+?)\s+(\d+)\s*$", line, re.I)
                if m2:
                    chap = int(m2.group(1))
                    title = m2.group(2).strip(" -")
                    if title and chap not in seen:
                        seen.add(chap)
                        out.append((chap, title))
                else:
                    # Continuation lines on some TOC pages.
                    if pending_chap is not None:
                        extra = re.sub(r"\.{3,}\s*\d+\s*$", "", line).strip(" -")
                        if extra:
                            pending_title = f"{pending_title} {extra}".strip()
                    continue
                continue

            chap = int(m.group(1))
            title = m.group(2).strip(" -")
            if not title:
                pending_chap = chap
                pending_title = ""
                continue
            if chap in seen:
                pending_chap = None
                pending_title = ""
                continue
            # If line likely wraps and has no trailing page number, store pending.
            if not re.search(r"\d+\s*$", line):
                pending_chap = chap
                pending_title = title
                continue
            seen.add(chap)
            out.append((chap, title))
            pending_chap = None
            pending_title = ""

        if pending_chap is not None and pending_title and pending_chap not in seen:
            seen.add(pending_chap)
            out.append((pending_chap, pending_title))
    out.sort(key=lambda x: x[0])
    return out


def linkify_html_text(
    text: str,
    current_page: int,
    page_file_map: Dict[int, str],
    figure_table_targets: Dict[str, Tuple[int, str]],
    chapter_targets: Dict[str, Tuple[int, str]],
) -> str:
    def repl(m: re.Match[str]) -> str:
        label = m.group(0)
        key = label.lower()
        target = figure_table_targets.get(key) or chapter_targets.get(key)
        if not target:
            return label
        t_page, t_anchor = target
        href = f"#{t_anchor}" if t_page == current_page else f"{page_file_map[t_page]}#{t_anchor}"
        return f'<a href="{escape_attr(href)}">{label}</a>'

    # Avoid linkifying inside existing tags by splitting tags/text parts.
    parts = re.split(r"(<[^>]+>)", text)
    for i, p in enumerate(parts):
        if p.startswith("<") and p.endswith(">"):
            continue
        parts[i] = REF_RE.sub(repl, p)
    return "".join(parts)


def html_to_plain_text(s: str) -> str:
    t = re.sub(r"<[^>]+>", "", s)
    t = html.unescape(t)
    t = t.replace("\ufffd", "")
    t = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", " ", t)
    t = re.sub(r"[\u200B-\u200D\uFEFF]", "", t)
    t = re.sub(r"\s+", " ", t).strip()
    return t


def page_has_contents_heading(blocks: Sequence[ContentBlock]) -> bool:
    for b in blocks:
        if b.kind not in {"h1", "h2", "p"}:
            continue
        t = html_to_plain_text(b.text).upper()
        if "CONTENTS IN BRIEF" in t:
            return True
    return False


def page_has_main_contents_heading(blocks: Sequence[ContentBlock]) -> bool:
    for b in blocks:
        if b.kind not in {"h1", "h2", "p"}:
            continue
        t = html_to_plain_text(b.text).upper()
        if t == "CONTENTS":
            return True
    return False


def page_has_heading_text(blocks: Sequence[ContentBlock], target: str) -> bool:
    t0 = target.upper().strip()
    for b in blocks:
        if b.kind not in {"h1", "h2", "p"}:
            continue
        t = html_to_plain_text(b.text).upper()
        if t == t0:
            return True
    return False


def extract_outline_contents_entries(doc: fitz.Document) -> List[Tuple[int, str, int]]:
    """
    Build clean contents entries from PDF outline:
    - level 1: chapter titles (and key front matter lines before chapters)
    - level 2: section titles under chapters
    """
    out: List[Tuple[int, str, int]] = []
    rows = doc.get_toc() or []
    started_chapters = False
    in_chapter = False
    for level, title, page in rows:
        t = normalize_text(str(title))
        if not t or page <= 0:
            continue
        if level == 1 and re.match(r"^Chapter\s+\d+\.", t, re.I):
            started_chapters = True
            in_chapter = True
            out.append((1, t, int(page)))
            continue
        if level == 2 and in_chapter:
            out.append((2, t, int(page)))
            continue
        if not started_chapters and level == 1:
            if t in {"List of Business Snapshots", "List of Technical Notes", "Preface"}:
                out.append((1, t, int(page)))
    return out


def split_contents_entries_for_pages(
    entries: Sequence[Tuple[int, str, int]], page_count: int
) -> List[List[Tuple[int, str, int]]]:
    if page_count <= 1:
        return [list(entries)]
    if not entries:
        return [[] for _ in range(page_count)]

    def weight(item: Tuple[int, str, int]) -> int:
        lvl, title, _ = item
        base = 12 if lvl == 1 else 9
        return 1 + (len(title) // base)

    total = sum(weight(e) for e in entries)
    target = max(1, total // page_count)
    chunks: List[List[Tuple[int, str, int]]] = []
    cur: List[Tuple[int, str, int]] = []
    cur_w = 0
    remain_chunks = page_count
    remain_entries = list(entries)
    while remain_chunks > 1 and remain_entries:
        nxt = remain_entries.pop(0)
        w = weight(nxt)
        if cur and cur_w + w > target:
            chunks.append(cur)
            cur = [nxt]
            cur_w = w
            remain_chunks -= 1
        else:
            cur.append(nxt)
            cur_w += w
    cur.extend(remain_entries)
    chunks.append(cur)
    while len(chunks) < page_count:
        chunks.append([])
    return chunks


def extract_business_snapshot_entries(doc: fitz.Document) -> List[Tuple[str, int]]:
    out: List[Tuple[str, int]] = []
    seen = set()
    for p in range(min(len(doc), 40)):
        txt = doc[p].get_text("text")
        if "BUSINESS SNAPSHOTS" not in txt.upper():
            continue
        for raw in txt.splitlines():
            line = normalize_text(raw)
            if not line:
                continue
            m = re.match(r"^(\d+\.\d+)\s+(.+?)\s*\.{2,}\s*(\d+)\s*$", line)
            if not m:
                continue
            label = f"{m.group(1)} {m.group(2)}"
            page = int(m.group(3))
            key = (m.group(1), page)
            if key in seen:
                continue
            seen.add(key)
            out.append((label, page))
    return out


def extract_technical_notes_entries(doc: fitz.Document) -> List[str]:
    out: List[str] = []
    for p in range(min(len(doc), 40)):
        txt = doc[p].get_text("text")
        if not re.search(r"(?m)^\s*TECHNICAL NOTES\s*$", txt):
            continue
        lines = [normalize_text(x) for x in txt.splitlines()]
        lines = [x for x in lines if x]
        # Keep only numbered entries; merge wrapped continuation lines.
        i = 0
        while i < len(lines):
            line = lines[i]
            if re.match(r"^\d+\.\s+", line):
                cur = line
                j = i + 1
                while j < len(lines) and not re.match(r"^\d+\.\s+", lines[j]):
                    # stop at known non-entry footer/header lines
                    if lines[j].upper() in {
                        "TECHNICAL NOTES",
                        "AVAILABLE ON THE AUTHOR'S WEBSITE",
                        "AVAILABLE ON THE AUTHOR’S WEBSITE",
                    }:
                        break
                    if lines[j].lower().startswith("www-2.rotman.utoronto.ca"):
                        break
                    if re.fullmatch(r"\d+", lines[j]):
                        break
                    cur = f"{cur} {lines[j]}"
                    j += 1
                out.append(cur)
                i = j
                continue
            i += 1
        break
    return out


def add_bookmark(paragraph, name: str, bid: int) -> None:
    if OxmlElement is None or qn is None:
        return
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
    if OxmlElement is None or qn is None:
        paragraph.add_run(text)
        return
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("w:anchor"), anchor)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hl.append(r)
    paragraph._p.append(hl)


def add_paragraph_with_links(paragraph, text: str, link_targets: Dict[str, str]) -> None:
    last = 0
    for m in REF_RE.finditer(text):
        if m.start() > last:
            paragraph.add_run(text[last:m.start()])
        label = m.group(0)
        anchor = link_targets.get(label.lower())
        if anchor:
            add_internal_hyperlink(paragraph, label, anchor)
        else:
            paragraph.add_run(label)
        last = m.end()
    if last < len(text):
        paragraph.add_run(text[last:])


def write_docx_from_blocks(
    chapter_title: str,
    blocks: List[ContentBlock],
    img_dir: Path,
    out_docx: Path,
) -> Dict[str, int]:
    if DocxDocument is None or WD_ALIGN_PARAGRAPH is None:
        raise RuntimeError("Missing dependency 'python-docx'. Install with: python3 -m pip install python-docx")

    doc = DocxDocument()
    section = doc.sections[0]
    usable_width = section.page_width - section.left_margin - section.right_margin

    # Assign and collect anchors first.
    fig_table_targets, chapter_targets = build_reference_maps(blocks)
    docx_link_targets: Dict[str, str] = {}
    for k, (_, anchor) in fig_table_targets.items():
        docx_link_targets[k] = slugify(anchor)
    for k, (_, anchor) in chapter_targets.items():
        docx_link_targets[k] = slugify(anchor)

    title = doc.add_heading(html_to_plain_text(chapter_title), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    fig_count = 0
    para_count = 0
    eq_count = 0
    h2_count = 0
    bookmark_id = 1

    for b in blocks:
        if b.kind == "h1":
            p = doc.add_heading(html_to_plain_text(b.text), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if b.anchor_id:
                add_bookmark(p, slugify(b.anchor_id), bookmark_id)
                bookmark_id += 1
            continue

        if b.kind == "h2":
            p = doc.add_heading(html_to_plain_text(b.text), level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if b.anchor_id:
                add_bookmark(p, slugify(b.anchor_id), bookmark_id)
                bookmark_id += 1
            h2_count += 1
            continue

        if b.kind == "equation":
            p = doc.add_paragraph(html_to_plain_text(b.text))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            eq_count += 1
            continue

        if b.kind == "figure":
            cap = html_to_plain_text(b.text)
            if cap:
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                if b.fig_id:
                    add_bookmark(cp, slugify(b.fig_id), bookmark_id)
                    bookmark_id += 1
                add_paragraph_with_links(cp, cap, docx_link_targets)
            if b.img_name:
                img_path = img_dir / Path(b.img_name).name
                if img_path.exists():
                    ip = doc.add_paragraph()
                    ip.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    ip.add_run().add_picture(img_path.as_posix(), width=usable_width)
                    fig_count += 1
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_paragraph_with_links(p, html_to_plain_text(b.text), docx_link_targets)
        para_count += 1

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx.as_posix())
    return {
        "paragraphs": para_count,
        "headings": h2_count,
        "equations": eq_count,
        "figures": fig_count,
    }


def score_validation(stats: Dict[str, int]) -> int:
    # Lower score is better.
    return (
        stats.get("long_paragraphs", 0) * 8
        + max(0, stats.get("captions", 0) - stats.get("figures", 0)) * 10
        + max(0, 3 - stats.get("figure_suppressed_lines", 0)) * 2
    )


def build_once(
    doc: fitz.Document,
    start_page: int,
    end_page: int,
    chapter_id: str,
    cfg: Dict[str, object],
    img_dir: Path,
) -> BuildResult:
    all_blocks: List[ContentBlock] = []
    fig_counter = 0
    page_stats: Dict[int, Dict[str, int]] = {}
    agg = {
        "captions": 0,
        "figures": 0,
        "figure_suppressed_lines": 0,
        "long_paragraphs": 0,
        "style_runs_italic": 0,
        "style_runs_bold": 0,
    }

    for page_num in range(start_page, end_page + 1):
        page = doc[page_num - 1]
        lines, visuals = extract_page_lines(page)
        blocks, fig_counter, st = compose_content_blocks_for_page(
            lines=lines,
            visuals=visuals,
            page=page,
            page_num=page_num,
            chapter_id=chapter_id,
            cfg=cfg,
            img_dir=img_dir,
            figure_start_idx=fig_counter,
        )
        all_blocks.extend(blocks)
        kinds = {"p": 0, "h1": 0, "h2": 0, "equation": 0, "figure": 0}
        for b in blocks:
            kinds[b.kind] = kinds.get(b.kind, 0) + 1
        page_stats[page_num] = kinds
        for k, v in st.items():
            agg[k] = agg.get(k, 0) + v

    score = score_validation(agg)
    agg["score"] = score
    return BuildResult(
        blocks=all_blocks,
        images_written=agg.get("figures", 0),
        validation=agg,
        page_stats=page_stats,
    )


def convert(
    pdf_path: Path,
    start_page: int,
    end_page: int,
    chapter_id: str,
    chapter_title: str,
    out_epub: Path,
    out_docx: Optional[Path],
    cfg: Dict[str, object],
) -> Dict[str, object]:
    if start_page < 1 or end_page < start_page:
        raise ValueError("Invalid page range")
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    out_epub.parent.mkdir(parents=True, exist_ok=True)
    report_path = out_epub.with_suffix(".validation.json")

    tmp_dir = Path(tempfile.mkdtemp(prefix="pdf2epub_"))
    try:
        doc = fitz.open(pdf_path.as_posix())
        if end_page > len(doc):
            raise ValueError(f"End page {end_page} exceeds document size {len(doc)}")

        # Two attempts: default and figure-expansion variant. Keep better score.
        attempts: List[Tuple[str, Dict[str, object]]] = []
        attempts.append(("default", dict(cfg)))
        cfg_wide = dict(cfg)
        cfg_wide["thresholds.figure_y_extension_ratio"] = min(
            0.46, float(get_opt(cfg, "thresholds.figure_y_extension_ratio")) + 0.1
        )
        cfg_wide["thresholds.figure_x_margin_ratio"] = max(
            0.03, float(get_opt(cfg, "thresholds.figure_x_margin_ratio")) - 0.02
        )
        attempts.append(("wide_figure", cfg_wide))

        best: Optional[BuildResult] = None
        best_name = ""
        best_img_dir: Optional[Path] = None

        run_dirs: List[Path] = []
        for name, cfg_attempt in attempts:
            run_dir = tmp_dir / f"run_{name}"
            run_img_dir = run_dir / "images"
            run_img_dir.mkdir(parents=True, exist_ok=True)
            run_dirs.append(run_dir)

            result = build_once(
                doc=doc,
                start_page=start_page,
                end_page=end_page,
                chapter_id=chapter_id,
                cfg=cfg_attempt,
                img_dir=run_img_dir,
            )

            if best is None or int(result.validation["score"]) < int(best.validation["score"]):
                best = result
                best_name = name
                best_img_dir = run_img_dir

        assert best is not None and best_img_dir is not None

        css = build_css()
        official_toc = extract_contents_chapter_toc(doc)

        book = epub.EpubBook()
        # Use a stable identifier to avoid creating duplicate entries in Books.app.
        build_tag = str(get_opt(cfg, "epub_build_tag"))
        book.set_identifier(f"pdf2epub-{chapter_id}-{start_page}-{end_page}-{build_tag}")
        book.set_title(chapter_title)
        book.set_language("en")
        book.add_author("John C. Hull")
        book.add_metadata("DC", "description", "Hybrid EPUB generated from PDF")

        css_item = epub.EpubItem(
            uid="style_main",
            file_name="style.css",
            media_type="text/css",
            content=css.encode("utf-8"),
        )
        book.add_item(css_item)

        # Build one EPUB section per source PDF page to allow page-by-page validation in Books.
        by_page: Dict[int, List[ContentBlock]] = {}
        for b in best.blocks:
            by_page.setdefault(b.page_num, []).append(b)
        page_file_map: Dict[int, str] = {p: f"{chapter_id}_p{p}.xhtml" for p in by_page.keys()}
        figure_table_targets, chapter_targets = build_reference_maps(best.blocks)
        outline_entries = extract_outline_contents_entries(doc)
        snapshot_entries = extract_business_snapshot_entries(doc)
        technical_entries = extract_technical_notes_entries(doc)
        main_contents_pages = sorted(p for p in by_page.keys() if page_has_main_contents_heading(by_page[p]))
        snapshot_pages = sorted(p for p in by_page.keys() if page_has_heading_text(by_page[p], "BUSINESS SNAPSHOTS"))
        technical_pages = sorted(p for p in by_page.keys() if page_has_heading_text(by_page[p], "TECHNICAL NOTES"))
        main_contents_chunks: Dict[int, List[Tuple[int, str, str]]] = {}
        snapshot_chunks: Dict[int, List[Tuple[str, str]]] = {}
        technical_chunks: Dict[int, List[str]] = {}
        if outline_entries and main_contents_pages:
            chunks = split_contents_entries_for_pages(outline_entries, len(main_contents_pages))
            for pnum, chunk in zip(main_contents_pages, chunks):
                rendered: List[Tuple[int, str, str]] = []
                for lvl, label, t_page in chunk:
                    if t_page in page_file_map:
                        href = page_file_map[t_page]
                    else:
                        href = "#"
                    rendered.append((lvl, label, href))
                main_contents_chunks[pnum] = rendered
        if snapshot_entries and snapshot_pages:
            chunks = split_contents_entries_for_pages([(1, a, b) for a, b in snapshot_entries], len(snapshot_pages))
            for pnum, chunk in zip(snapshot_pages, chunks):
                rendered: List[Tuple[str, str]] = []
                for _lvl, label, t_page in chunk:
                    href = page_file_map[t_page] if t_page in page_file_map else "#"
                    rendered.append((label, href))
                snapshot_chunks[pnum] = rendered
        if technical_entries and technical_pages:
            chunks = split_contents_entries_for_pages([(1, x, 0) for x in technical_entries], len(technical_pages))
            for pnum, chunk in zip(technical_pages, chunks):
                technical_chunks[pnum] = [label for _lvl, label, _tp in chunk]

        page_items: List[epub.EpubHtml] = []
        for pnum in sorted(by_page.keys()):
            title = f"{chapter_title} - p{pnum}"
            toc_override: Optional[List[Tuple[str, str]]] = None
            toc_main_override: Optional[List[Tuple[int, str, str]]] = None
            business_snapshots_override: Optional[List[Tuple[str, str]]] = None
            technical_notes_override: Optional[List[str]] = None
            if official_toc and page_has_contents_heading(by_page[pnum]):
                toc_override = []
                for chap_no, chap_title in official_toc:
                    t = chapter_targets.get(f"chapter {chap_no}")
                    if t:
                        t_page, t_anchor = t
                        href = (
                            f"#{t_anchor}"
                            if t_page == pnum
                            else f"{page_file_map[t_page]}#{t_anchor}"
                        )
                    else:
                        href = "#"
                    toc_override.append((f"{chap_no}. {chap_title}", href))
            if pnum in main_contents_chunks:
                toc_main_override = main_contents_chunks[pnum]
            if pnum in snapshot_chunks:
                business_snapshots_override = snapshot_chunks[pnum]
            if pnum in technical_chunks:
                technical_notes_override = technical_chunks[pnum]
            xhtml = chapter_to_xhtml(
                chapter_title=chapter_title,
                blocks=by_page[pnum],
                current_page=pnum,
                page_file_map=page_file_map,
                figure_table_targets=figure_table_targets,
                chapter_targets=chapter_targets,
                toc_override=toc_override,
                toc_main_override=toc_main_override,
                business_snapshots_override=business_snapshots_override,
                technical_notes_override=technical_notes_override,
            )
            item = epub.EpubHtml(
                title=title,
                file_name=page_file_map[pnum],
                lang="en",
            )
            item.set_content(xhtml)
            book.add_item(item)
            page_items.append(item)

        image_qc: Dict[str, object] = {"total_images": 0, "trailing_after_bar": 0, "flagged": []}
        for img in sorted(best_img_dir.glob("*.png")):
            trim_image_whitespace(img)
            trail = gray_bar_trailing_pixels(img)
            image_qc["total_images"] = int(image_qc["total_images"]) + 1
            if trail > 0:
                image_qc["trailing_after_bar"] = int(image_qc["trailing_after_bar"]) + 1
                image_qc["flagged"].append({"image": img.name, "trailing_pixels": trail})
            book.add_item(
                epub.EpubImage(
                    uid=f"img_{img.stem}",
                    file_name=f"images/{img.name}",
                    media_type="image/png",
                    content=img.read_bytes(),
                )
            )

        toc_entries: List[object] = []
        if official_toc:
            for chap_no, chap_title in official_toc:
                t = chapter_targets.get(f"chapter {chap_no}")
                if not t:
                    continue
                pnum, anchor = t
                href = f"{page_file_map[pnum]}#{anchor}"
                uid = f"toc-ch{chap_no}"
                toc_entries.append(epub.Link(href, f"{chap_no}. {chap_title}", uid))
        else:
            noisy_toc_pages = {
                p for p, st in best.page_stats.items() if int(st.get("h2", 0)) >= 8 and int(st.get("p", 0)) <= 6
            }
            for b in best.blocks:
                if b.kind != "h2":
                    continue
                if b.page_num in noisy_toc_pages:
                    continue
                if not b.anchor_id:
                    continue
                text = html_to_plain_text(b.text)
                if not SECTION_RE.match(text):
                    continue
                if re.search(r"\.{3,}|…{2,}", text):
                    continue
                if re.search(r"\d{2,}\s*$", text):
                    continue
                if len(text) < 2:
                    continue
                href = f"{page_file_map[b.page_num]}#{b.anchor_id}"
                toc_entries.append(epub.Link(href, text, b.anchor_id))
        book.toc = tuple(toc_entries if toc_entries else page_items)
        book.spine = page_items + ["nav"]
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        epub.write_epub(out_epub.as_posix(), book, {})

        docx_stats: Optional[Dict[str, int]] = None
        if out_docx is not None:
            docx_stats = write_docx_from_blocks(
                chapter_title=chapter_title,
                blocks=best.blocks,
                img_dir=best_img_dir,
                out_docx=out_docx,
            )

        report = {
            "selected_attempt": best_name,
            "page_range": [start_page, end_page],
            "validation": best.validation,
            "images_written": best.images_written,
            "page_stats": best.page_stats,
            "image_qc": image_qc,
            "output": out_epub.as_posix(),
            "docx_output": out_docx.as_posix() if out_docx is not None else None,
            "docx_stats": docx_stats,
        }
        report_path.write_text(json.dumps(report, indent=2), encoding="utf-8")
        return report

    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Convert PDF chapter range to hybrid EPUB")
    p.add_argument("--pdf", required=False, help="Path to source PDF")
    p.add_argument("--start", required=False, type=int, help="Start page (1-based)")
    p.add_argument("--end", required=False, type=int, help="End page (1-based)")
    p.add_argument("--chapter-id", required=False, help="Chapter id, e.g. ch10")
    p.add_argument("--title", required=False, help="Chapter title")
    p.add_argument("--out", required=False, help="Output EPUB path")
    p.add_argument("--out-docx", required=False, help="Output DOCX path (Pages-openable)")
    p.add_argument("--config", default="config/ch10.yaml", help="Config file path")
    return p.parse_args()


def main() -> None:
    args = parse_args()
    cfg = load_simple_yaml(Path(args.config))

    pdf = Path(args.pdf or str(cfg.get("pdf_path", ""))).expanduser()
    start = int(args.start if args.start is not None else cfg.get("start_page", 0))
    end = int(args.end if args.end is not None else cfg.get("end_page", 0))
    chapter_id = str(args.chapter_id or cfg.get("chapter_id", "chapter"))
    title = str(args.title or cfg.get("chapter_title", "Chapter"))
    out = Path(args.out or str(cfg.get("out_path", "output/chapter.epub"))).expanduser()
    out_docx = Path(args.out_docx).expanduser() if args.out_docx else out.with_suffix(".docx")

    if not pdf:
        raise SystemExit("Missing PDF path")
    if start <= 0 or end <= 0:
        raise SystemExit("Missing valid page range")

    report = convert(
        pdf_path=pdf,
        start_page=start,
        end_page=end,
        chapter_id=chapter_id,
        chapter_title=title,
        out_epub=out,
        out_docx=out_docx,
        cfg=cfg,
    )
    print(f"Done: {out}")
    print(f"DOCX: {out_docx}")
    print(f"Validation: {json.dumps(report['validation'])}")


if __name__ == "__main__":
    main()
